"""
File Reader — Extract text dari file yang diupload.

Pipeline:
  Upload → Validasi Format & Ukuran → Extract Text → Clean Text → …
"""

import io
from fastapi import UploadFile, HTTPException
from docx import Document
import pdfplumber

# Format file yang didukung
ALLOWED_EXTENSIONS = {".txt", ".pdf", ".docx"}
ALLOWED_CONTENT_TYPES = {
    "text/plain",
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

# Batas ukuran file: 20MB
MAX_FILE_SIZE_BYTES = 20 * 1024 * 1024


def validate_file(file: UploadFile, content: bytes) -> str:
    """
    Validasi format, content-type, dan ukuran file.
    Returns extension jika valid, raise HTTPException jika tidak.
    """
    # Validasi ukuran
    if len(content) > MAX_FILE_SIZE_BYTES:
        raise HTTPException(
            status_code=413,
            detail=f"Ukuran file melebihi batas maksimum "
                   f"{MAX_FILE_SIZE_BYTES // (1024 * 1024)}MB.",
        )

    # Validasi ekstensi
    filename = file.filename or ""
    ext = ("." + filename.rsplit(".", 1)[-1].lower()) if "." in filename else ""
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Format '{ext or 'unknown'}' tidak didukung. "
                   f"Gunakan: {', '.join(sorted(ALLOWED_EXTENSIONS))}",
        )

    # Validasi content-type
    content_type = file.content_type or ""
    if content_type not in ALLOWED_CONTENT_TYPES:
        raise HTTPException(
            status_code=400,
            detail=f"Content-type '{content_type}' tidak diizinkan.",
        )

    return ext


async def read_file(file: UploadFile) -> str:
    """
    Extract text dari file. Validasi dilakukan di dalam fungsi ini.
    Mendukung: TXT, PDF (pdfplumber), DOCX.
    """
    # Baca konten sekali — setelah ini pointer file sudah di akhir
    content = await file.read()

    # Validasi format, content-type, dan ukuran
    ext = validate_file(file, content)

    # ---- TXT ----
    if ext == ".txt":
        return content.decode("utf-8", errors="ignore").strip()

    # ---- DOCX ----
    elif ext == ".docx":
        doc = Document(io.BytesIO(content))

        parts = []

        # Ambil paragraf biasa
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())

        # Ambil teks dari tabel (doc.paragraphs tidak include teks di tabel)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())

        return "\n".join(parts)

    # ---- PDF (pdfplumber) ----
    elif ext == ".pdf":
        parts = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text and text.strip():
                    parts.append(text.strip())

        return "\n".join(parts)